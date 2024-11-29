#coding=utf-8
import os

from docx import Document
from win32com import client
from docxcompose.composer import Composer
from pdf2docx import parse

class WordConvertor:
    # wdFormatPDF = 17
    # wdFormatDoc = 0
    # wdFormatDocx = 12    #
    #

    def __init__(self):
        self._doc_app = client.Dispatch('Word.Application')
        self._documents = self._doc_app.Documents

    @property
    def docFormat(self):
        wdFormat = {'doc':0,
                    'docx':12,
                    'pdf':17
        }
        return wdFormat

    def doc_to_docx(self, doc_path, docx_path=None):
        if doc_path.endswith('.doc'):
            if docx_path is None:
                docx_path = "{}x".format(doc_path)
            doc = self._documents.Open(doc_path)
            doc.SaveAs(docx_path, self.docFormat['docx'])
            doc.Close()
        return docx_path

    def doc_to_pdf(self, doc_path, pdf_path=None):
        if doc_path.endswith('.doc') or doc_path.endswith('.docx'):
            if pdf_path is None:
                pdf_path = os.path.splitext(doc_path)[0] + '.pdf'
                doc = self._documents.Open(doc_path)
                doc.SaveAs(pdf_path, self.docFormat['pdf'])
                doc.Close()
        return pdf_path

    def pdf_to_doc(self, pdf_path, doc_path=None):
        if doc_path is None:
            doc_path = os.path.splitext(pdf_path)[0] + '.docx'
        parse(pdf_path,doc_path)
        return doc_path
    def __del__(self):
        pass
        #self._doc_app.Quit()



class WordObject:
    # 导入 docx 文件数据

    def __init__(self, file_path=None):
        self._doc_path = file_path
        self.load_doc(file_path)


    def print_hi(self):
        print(type(self._doc))

    @property
    def doc(self):
        return self._doc


    def load_doc(self, doc_path):
        self._doc_path = doc_path
        self._doc = Document(doc_path)


    @property
    def doc_path(self):
        return self._doc_path

    @doc_path.setter
    def doc_path(self, value):
        self._doc_path = value
        self.load_doc(self._doc_path)

    def insert_txt(self, txt, font_name='宋体'):
        para = self._doc.add_paragraph(txt)
        for run in para.runs:
            font = run.font
            font.name = font_name

    def insert_txt_file(self, txt_path, encoding='utf-8', font_name='宋体'):
        with open(txt_path, 'r', encoding=encoding) as file:
            txt_lines = file.readlines()
        for txt_line in txt_lines:
            para = self._doc.add_paragraph(txt_line)
            for run in para.runs:
                font = run.font
                font.name = font_name

    def insert_img(self, img_path):
        section = self._doc.add_section()
        img = self._doc.add_picture(img_path)
        page_width = section.page_width - section.left_margin - section.right_margin
        page_height = section.page_height - section.top_margin - section.bottom_margin
        rw = img.width / page_width
        rh = img.height / page_height
        if (rw < 1) and (rh < 1):
            pass
        elif rw >= rh:
            img.width = int(img.width / rw)
            img.height = int(img.height / rw)
        else:
            img.width = int(img.width / rh)
            img.height = int(img.height / rh)
        #self._doc.add_page_break()

    def insert_docx_org(self, doc_path):
        master = self._doc
        composer = Composer(master)
        print(doc_path)
        new_doc = WordObject(doc_path)
        try:
            composer.append(new_doc.doc)
            composer.save(self._doc_path)
            self.load_doc(self._doc_path)
        except Exception as err:
            print("[Error in Composing]:" + str(err))

    def insert_docx(self, doc_path):
        try:
            new_doc = Document(doc_path)
            for element in new_doc.element.body:
                self._doc.element.body.append(element)
            return ""
        except Exception as err:
            print("[Error in Composing]:" + str(err))
            return str(err)

    def delete_blank_pages(self):
        for sec in self._doc.sections:
            #head = sec.header
            #foot = sec.footer
            page = sec._element.getnext()
            if page is None:
                continue
            for element in page.iter():
                if element.text.strip():
                    break
                else:
                    sec.getparent().remove(sec)
                    continue


        #self.save()


    def save(self, file_path = None):
        if file_path is None:
            file_path = self._doc_path
        self._doc.save(file_path)
        self._doc_path = file_path



